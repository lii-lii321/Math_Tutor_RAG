"""
My Progress 页面视图
显示错题本和进度管理
"""
import streamlit as st
import os
import time
import streamlit_antd_components as sac
from db_manager import DBManager
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def generate_word_exam(questions, exam_title="错题复习试卷"):
    """生成 Word 试卷"""
    doc = Document()
    heading = doc.add_heading(exam_title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"【第 {idx} 题】").bold = True
        if os.path.exists(q['image_path']):
            try: 
                doc.add_picture(q['image_path'], width=Inches(4.5))
            except: 
                pass
        doc.add_paragraph("\n" * 2)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def render_progress_page(user):
    """渲染 My Progress 页面"""
    st.markdown("### 📒 My Progress")
    
    db = DBManager()
    full_history = db.get_history(user['id'], user['role'])
    
    default_search = st.session_state.get('search_query', "")
    if default_search: 
        st.session_state['search_query'] = None
    
    st.markdown("""
    <div class="dashboard-card" style="margin-bottom: 1rem;">
    """, unsafe_allow_html=True)
    
    c_s, c_r, c_ex = st.columns([3, 1, 1.5])
    with c_s: 
        search_term = st.text_input("🔍 Search...", value=default_search, key="unique_search_bar")
    with c_r: 
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄 Refresh", use_container_width=True): 
            st.rerun()
    
    history = [item for item in full_history 
               if search_term.lower() in item['tags'].lower() or search_term.lower() in item['ai_content'].lower()] \
               if search_term else full_history
    
    with c_ex:
        st.markdown("<br>", unsafe_allow_html=True)
        if history:
            doc_io = generate_word_exam(history, f"MathMaster - {search_term if search_term else 'All'} Review")
            st.download_button(
                "📥 Export", 
                data=doc_io, 
                file_name=f"错题卷_{int(time.time())}.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                type="primary", 
                use_container_width=True
            )
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    if not history:
        sac.result(label='No questions yet', status='empty')
    else:
        st.caption(f"Total: {len(history)} questions")
        
        selected_ids = []
        for item in history:
            with st.expander(f"🏷️ {item['tags']} | 📅 {item['date']} | 🆔 {item['id']}"):
                c_info, c_del = st.columns([6, 1])
                c_info.info(f"Date: {item['date']} | User: {item['username']}")
                if c_del.button("🗑️ Delete", key=f"del_{item['id']}", type="secondary"):
                    db.delete_question(item['id'])
                    st.rerun()
                st.divider()
                c_img, c_edit = st.columns([1, 2])
                if os.path.exists(item['image_path']): 
                    c_img.image(item['image_path'], use_container_width=True)
                else: 
                    c_img.error("❌ Image missing")
                
                t_view, t_edit = c_edit.tabs(["👀 Preview", "✏️ Edit"])
                t_view.markdown(
                    f"<div style='background:#f8f9fa;padding:15px;border-radius:10px;'>{item['ai_content']}</div>", 
                    unsafe_allow_html=True
                )
                with t_edit.form(f"edit_{item['id']}"):
                    n_tags = st.text_input("Tags", value=item['tags'])
                    n_con = st.text_area("Analysis", value=item['ai_content'], height=300)
                    if st.form_submit_button("💾 Save"):
                        db.update_question(item['id'], n_con, n_tags)
                        st.rerun()
                
                if st.checkbox("Select", key=f"chk_{item['id']}"):
                    selected_ids.append(item['id'])
        
        if selected_ids:
            st.warning(f"Selected {len(selected_ids)} questions")
            if st.button(f"🗑️ Batch Delete ({len(selected_ids)})", type="primary"):
                for qid in selected_ids: 
                    db.delete_question(qid)
                st.rerun()
