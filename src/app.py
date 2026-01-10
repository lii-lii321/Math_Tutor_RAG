import streamlit as st
import base64
import os
import time
import io
import concurrent.futures
import threading
from PIL import Image
from dotenv import load_dotenv
import streamlit_antd_components as sac
from db_manager import DBManager
from streamlit.runtime.scriptrunner import add_script_run_ctx, get_script_run_ctx
from streamlit_echarts import st_echarts
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

# ================= 🟢 1. 关键修复：网络代理配置 =================
# 如果你在国内使用 VPN，通常需要手动告诉 Python 走这个端口
# 常见的端口是 7890 (Clash) 或 10809 (v2ray)
# Clash Verge 默认端口通常是 7897，我们试一下！
os.environ["HTTP_PROXY"] = "http://127.0.0.1:7890"
os.environ["HTTPS_PROXY"] = "http://127.0.0.1:7890"

# ================= 2. 全局配置与样式 =================

def load_css():
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700&display=swap');
        html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3436; }
        .stApp { background-color: #FDFBF7; }
        [data-testid="stSidebar"] { background-color: #FFFFFF; box-shadow: 2px 0 20px rgba(0,0,0,0.02); }
        .cream-card {
            background-color: #FFFFFF; border-radius: 24px; padding: 30px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.05); margin-bottom: 25px;
        }
        .stButton>button { border-radius: 12px; font-weight: bold; border: none; transition: all 0.3s; width: 100%; }
        .stTextInput>div>div>input { border-radius: 12px; border: 2px solid #F0F2F5; background-color: #F9FAFB; }
    </style>
    """, unsafe_allow_html=True)

# ================= 3. 模型与工具函数 =================

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-flash-latest')
else:
    st.error("⚠️ 未找到 GOOGLE_API_KEY，请在 .env 文件中配置！")

DATA_DIR = "../data/full_page_book"
IMG_DIR = os.path.join(DATA_DIR, "images")
os.makedirs(IMG_DIR, exist_ok=True)
MAX_WORKERS = 1

st.set_page_config(page_title="MathMaster Edu", page_icon="✏️", layout="wide")

def process_single_file(file_obj, user_tags, hint, user_id, ctx):
    if ctx: add_script_run_ctx(threading.current_thread(), ctx)
    time.sleep(0.5) 
    fname = file_obj.name
    
    try:
        image = Image.open(file_obj)
        prompt = f"""
        你是一名亲切的小学数学老师。请对这张错题进行温柔、详细的讲解。
        【用户提示】：{hint if hint else "无"}
        
        请严格按照以下 Markdown 格式输出：
        ## 🧠 考点在哪里？
        (简要分析考点)
        ## 📝 老师来细讲
        (详细的步骤解析)
        ## ✅ 正确答案
        (给出最终结果)
        ## 🏷️ 标签
        (请提取 2-3 个核心知识点关键词，用逗号分隔。例如：几何, 相似三角形, 计算)
        """
        
        ai_content = ""
        last_error = ""
        
        # 尝试 3 次调用
        for i in range(3):
            try:
                response = model.generate_content([prompt, image])
                ai_content = response.text
                break # 成功就跳出
            except Exception as e:
                last_error = str(e)
                print(f"❌ Gemini API Error (Attempt {i+1}): {e}")
                time.sleep(2)

        # 🟢 修复逻辑：如果循环结束还没有内容，说明真的失败了
        if not ai_content:
            return False, fname, f"AI 连接失败: {last_error}", ""

        # 提取标签
        final_tags = user_tags
        match = re.search(r"## 🏷️ 标签[:：]?\s*(.*)", ai_content, re.DOTALL)
        if match:
            ai_extracted_tags = match.group(1).strip()
            ai_extracted_tags = ai_extracted_tags.replace("。", "").replace(".", "").strip()
            if ai_extracted_tags:
                final_tags = f"{user_tags}, {ai_extracted_tags}"
                tag_list = [t.strip() for t in final_tags.replace("，", ",").split(",") if t.strip()]
                final_tags = ", ".join(list(set(tag_list)))

        timestamp = str(int(time.time() * 1000))
        save_name = f"User{user_id}_{timestamp}.jpg"
        save_path = os.path.join(IMG_DIR, save_name)
        
        file_obj.seek(0)
        with open(save_path, "wb") as f:
            f.write(file_obj.read())
            
        db = DBManager()
        db.save_question(user_id, fname, ai_content, save_path, final_tags)
        return True, fname, ai_content, save_path
        
    except Exception as e:
        return False, fname, f"系统错误: {str(e)}", ""

# (generate_word_exam 函数保持不变，为了节省篇幅省略，不需要改动它)
def generate_word_exam(questions, exam_title="错题复习试卷"):
    doc = Document()
    heading = doc.add_heading(exam_title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"【第 {idx} 题】").bold = True
        if os.path.exists(q['image_path']):
            try: doc.add_picture(q['image_path'], width=Inches(4.5))
            except: pass
        doc.add_paragraph("\n" * 2)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ================= 4. 页面逻辑 =================

def show_login_page():
    st.markdown("<br><br><br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1.2, 1])
    with col2:
        st.markdown('<div class="cream-card"><h1 style="text-align: center;">✨ MathMaster</h1><p style="text-align: center; color: #B2BEC3;">让数学学习变得像呼吸一样简单</p></div>', unsafe_allow_html=True)
        with st.form("login_form"):
            username = st.text_input("账号", placeholder="Student ID")
            password = st.text_input("密码", type="password", placeholder="Password")
            if st.form_submit_button("🎈 登录系统", use_container_width=True):
                db = DBManager()
                user = db.login(username, password)
                if user:
                    st.session_state['user_info'] = {'id': user[0], 'username': username, 'role': user[1]}
                    st.rerun()
                else:
                    st.error("账号或密码错误")

def show_main_page():
    load_css()
    user = st.session_state['user_info']
    
    default_index = 0
    if st.session_state.get('navigate_to') == '错题本':
        default_index = 1
        st.session_state['navigate_to'] = None 
    
    with st.sidebar:
        st.markdown(f"""<div style="text-align: center; padding: 20px;"><div style="background: #FFEAA7; width: 60px; height: 60px; border-radius: 50%; line-height: 60px; font-size: 30px; margin: 0 auto;">🦁</div><h3>{user['username']}</h3><p>{user['role'].upper()}</p></div>""", unsafe_allow_html=True)
        menu_item = sac.menu([
            sac.MenuItem('学习中心', icon='book-half'),
            sac.MenuItem('错题本', icon='journal-bookmark-fill'),
            sac.MenuItem('设置', icon='gear-fill', type='group', children=[sac.MenuItem('退出登录', icon='box-arrow-right')]),
        ], index=default_index, format_func='title', color='orange', variant='light', open_all=True)

    if menu_item == '退出登录':
        st.session_state['user_info'] = None
        st.rerun()

    elif menu_item == '学习中心':
        st.title(f"早安, {user['username']}! ☀️")
        st.markdown('<div class="cream-card">', unsafe_allow_html=True)
        c1, c2 = st.columns([2, 1])
        with c1: uploaded_files = st.file_uploader("📥 上传作业图片", accept_multiple_files=True, type=['jpg','png'])
        with c2: 
            tags = st.text_input("🏷️ 本次标签", value="期末复习")
            hint = st.text_input("💡 小提示", placeholder="哪里不懂点哪里...")
            
        if uploaded_files:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🚀 开始魔法解析 (Gemini)", type="primary", use_container_width=True):
                ctx = get_script_run_ctx()
                progress = st.progress(0)
                status = st.status("🔮 AI 老师正在思考...", expanded=True)
                
                with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = []
                    for f in uploaded_files:
                        futures.append(executor.submit(process_single_file, f, tags, hint, user['id'], ctx))
                    
                    completed = 0
                    for future in concurrent.futures.as_completed(futures):
                        ok, fname, content, path = future.result()
                        completed += 1
                        progress.progress(completed / len(uploaded_files))
                        if ok:
                            status.write(f"✅ {fname} 完成")
                            with st.expander(f"📖 查看解析: {fname}"):
                                st.columns([1, 2])[0].image(path, use_container_width=True)
                                st.columns([1, 2])[1].markdown(content)
                        else:
                            # 🟢 修复逻辑：这里会显示具体的报错信息
                            status.error(f"❌ {fname} 失败: {content}")
                            st.error(f"⚠️ 错误详情: {content}")
                            
                status.update(label="🎉 队列处理结束", state="complete")
        st.markdown('</div>', unsafe_allow_html=True)

        # 🟢 4. 修复图表逻辑：解决文字重叠问题
        st.markdown("### 📊 学习状态分析")
        db = DBManager()
        all_history = db.get_history(user['id'], user['role'])
        if all_history:
            from collections import Counter
            all_tags = []
            for item in all_history:
                for t in item['tags'].replace('，', ',').split(','):
                    if t.strip(): all_tags.append(t.strip())
            
            top_tags = Counter(all_tags).most_common(5)
            c_chart, c_btn = st.columns([1.5, 2])
            with c_chart:
                # 🟢 ECharts 配置优化
                st_echarts(options={
                    "tooltip": {"trigger": "item"},
                    "legend": {
                        "bottom": "0%",   # 把图例移到底部
                        "left": "center",
                        "itemGap": 10     # 图例之间的间距
                    },
                    "series": [{
                        "name": "错题分布",
                        "type": "pie",
                        "radius": ["35%", "60%"], # 🟢 稍微调小一点，给文字留空间
                        "avoidLabelOverlap": True, # 🟢 防止重叠的魔法开关
                        "label": {
                            "show": True,
                            "position": "outside",
                            "formatter": "{b}: {c}" # 显示 标签: 数量
                        },
                        "data": [{"value": v, "name": k} for k, v in top_tags]
                    }]
                }, height="350px") # 🟢 增加高度
            with c_btn:
                st.caption("🔥 你的高频错题点")
                cols = st.columns(3)
                for idx, (t, c) in enumerate(top_tags):
                    if cols[idx % 3].button(f"{t}\n({c})", key=f"btn_{t}", use_container_width=True):
                        st.session_state['search_query'] = t
                        st.session_state['navigate_to'] = "错题本"
                        st.rerun()

    elif menu_item == '错题本':
        st.title("📒 我的错题本")
        default_search = st.session_state.get('search_query', "")
        if default_search: st.session_state['search_query'] = None 

        st.markdown('<div class="cream-card">', unsafe_allow_html=True)
        c_s, c_r, c_ex = st.columns([3, 1, 1.5]) 
        with c_s: search_term = st.text_input("🔍 搜索...", value=default_search, key="unique_search_bar")
        with c_r: 
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🔄 刷新", use_container_width=True): st.rerun()
            
        db = DBManager()
        full_history = db.get_history(user['id'], user['role'])
        history = [item for item in full_history if search_term in item['tags'] or search_term in item['ai_content']] if search_term else full_history

        with c_ex:
            st.markdown("<br>", unsafe_allow_html=True)
            if history:
                doc_io = generate_word_exam(history, f"MathMaster - {search_term if search_term else '综合'}复习")
                st.download_button("📥 导出试卷", data=doc_io, file_name=f"错题卷_{int(time.time())}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
            
        if not history:
            sac.result(label='空空如也', status='empty')
        else:
            c_all, _ = st.columns([1, 5])
            with c_all:
                if st.checkbox("全选", key="select_all"):
                    for item in history: st.session_state[f"chk_{item['id']}"] = True
            
            batch_ph = st.empty()
            st.caption(f"共 {len(history)} 题")
            
            selected_ids = []
            for item in history:
                c_chk, c_con = st.columns([0.05, 0.95])
                with c_chk:
                    st.write(""); st.write("")
                    if st.checkbox("", key=f"chk_{item['id']}"): selected_ids.append(item['id'])
                
                with c_con:
                    with st.expander(f"🏷️ {item['tags']} | 📅 {item['date']} | 🆔 {item['id']}"):
                        c_info, c_del = st.columns([6, 1])
                        c_info.info(f"录入: {item['date']} | 用户: {item['username']}")
                        if c_del.button("🗑️ 删除", key=f"del_{item['id']}", type="secondary"):
                            db.delete_question(item['id'])
                            st.rerun()
                        st.divider()
                        c_img, c_edit = st.columns([1, 2])
                        if os.path.exists(item['image_path']): c_img.image(item['image_path'], use_container_width=True)
                        else: c_img.error("❌ 图片丢失")
                        
                        t_view, t_edit = c_edit.tabs(["👀 预览", "✏️ 编辑"])
                        t_view.markdown(f"<div style='background:#f8f9fa;padding:15px;border-radius:10px;'>{item['ai_content']}</div>", unsafe_allow_html=True)
                        with t_edit.form(f"edit_{item['id']}"):
                            n_tags = st.text_input("标签", value=item['tags'])
                            n_con = st.text_area("解析", value=item['ai_content'], height=300)
                            if st.form_submit_button("💾 保存"):
                                db.update_question(item['id'], n_con, n_tags)
                                st.rerun()

            if selected_ids:
                with batch_ph.container():
                    st.warning(f"已选 {len(selected_ids)} 题")
                    if st.button(f"🗑️ 批量删除 ({len(selected_ids)})", type="primary"):
                        for qid in selected_ids: db.delete_question(qid)
                        st.session_state['select_all'] = False
                        st.rerun()

if 'user_info' not in st.session_state:
    st.session_state['user_info'] = None

if st.session_state['user_info'] is None:
    load_css()
    show_login_page()
else:
    show_main_page()