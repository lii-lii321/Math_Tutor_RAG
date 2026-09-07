"""导出服务：把错题列表排版为可打印的 Word 复习卷。"""
from __future__ import annotations

import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from backend.models.schemas import QuestionOut

_BLANK_LINES_AFTER_QUESTION = 4


def generate_word_exam(
    questions: list[QuestionOut], exam_title: str = "错题复习卷"
) -> io.BytesIO:
    doc = Document()
    heading = doc.add_heading(exam_title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"共 {len(questions)} 题 · 由 MathMaster Edu 自动生成 · "
        + dt_date_today()
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)

    for idx, question in enumerate(questions, 1):
        meta = doc.add_paragraph()
        run = meta.add_run(f"第 {idx} 题　[{question.difficulty}]　{' / '.join(question.tags)}")
        run.bold = True

        if question.image_path and os.path.exists(question.image_path):
            try:
                doc.add_picture(question.image_path, width=Inches(4.2))
            except Exception:  # noqa: BLE001 - 图片损坏不阻断导出
                doc.add_paragraph("(原图缺失)")

        doc.add_paragraph("\n" * _BLANK_LINES_AFTER_QUESTION)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def dt_date_today() -> str:
    import datetime as dt

    return dt.date.today().strftime("%Y-%m-%d")
